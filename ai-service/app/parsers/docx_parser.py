import io
from docx import Document
from .base import DocumentParser


class DocxParser(DocumentParser):
    def supports(self, file_type: str, file_name: str = "") -> bool:
        if file_type and (
            "wordprocessingml" in file_type.lower()
            or "docx" in file_type.lower()
            or "msword" in file_type.lower()
        ):
            return True
        if file_name and (file_name.lower().endswith(".docx") or file_name.lower().endswith(".doc")):
            return True
        return False

    def parse(self, content_bytes: bytes) -> str:
        doc = Document(io.BytesIO(content_bytes))
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        return "\n".join(text_parts).strip()
